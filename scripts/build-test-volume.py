"""
Semantic test dataset builder using FULL Wikipedia articles (NO CHUNKING).
Each document = one full Wikipedia article.
"""

import json
import time
import requests
from pathlib import Path
from tqdm import tqdm
from PIL import Image
from bs4 import BeautifulSoup
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document


# ============================================================
# Configuration
# ============================================================

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = PROJECT_ROOT / "data"
OUTPUT_DIR = PROJECT_ROOT / "semantic_test_dataset"

CATEGORY_MAPPING = {
    "boat": "boats",
    "airplane": "airplanes",
    "car": "cars",
    "dog": "animals"
}

IMAGES_PER_CATEGORY = 10


# ============================================================
# Wikipedia Sources (1 FILE = 1 ARTICLE)
# ============================================================

WIKIPEDIA_SOURCES = {
    "boats": [
        "Boat",
        "Ship",
        "Naval architecture",
        "Hull (watercraft)",
        "Sailing"
    ],
    "airplanes": [
        "Aircraft",
        "Airplane",
        "Aviation",
        "Jet engine",
        "Aerodynamics"
    ],
    "cars": [
        "Car",
        "Automobile",
        "Electric car",
        "Internal combustion engine",
        "Vehicle"
    ],
    "animals": [
        "Animal",
        "Dog",
        "Mammal",
        "Zoology",
        "Wildlife"
    ]
}


# ============================================================
# Verify COCO dataset exists
# ============================================================

def verify_coco():
    if not (DATA_DIR / "val2017").exists():
        raise RuntimeError("Missing data/val2017 folder.")
    if not (DATA_DIR / "annotations" / "instances_val2017.json").exists():
        raise RuntimeError("Missing instances_val2017.json.")
    print("[OK] COCO files verified.")


# ============================================================
# Create folder structure
# ============================================================

def create_structure():
    for folder in CATEGORY_MAPPING.values():
        (OUTPUT_DIR / folder / "images").mkdir(parents=True, exist_ok=True)
        (OUTPUT_DIR / folder / "texts").mkdir(parents=True, exist_ok=True)


# ============================================================
# Image processing
# ============================================================

def resize_and_save_image(src_path: Path, dst_base: Path):
    img = Image.open(src_path).convert("RGB")
    img.thumbnail((800, 800))

    img.save(dst_base.with_suffix(".jpg"), format="JPEG", quality=90)
    img.save(dst_base.with_suffix(".png"), format="PNG")


# ============================================================
# Wikipedia FULL article fetch
# ============================================================

def fetch_wikipedia_text(title: str):
    url = "https://en.wikipedia.org/w/api.php"

    headers = {
        "User-Agent": "semantic-retrieval-test/1.0"
    }

    params = {
        "action": "parse",
        "page": title,
        "prop": "text",
        "format": "json"
    }

    try:
        response = requests.get(url, params=params, headers=headers, timeout=10)
        response.raise_for_status()

        html = response.json()["parse"]["text"]["*"]
        text = clean_html(html)

        time.sleep(0.5)

        return text

    except Exception as e:
        print(f"[WARN] Failed to fetch {title}: {e}")
        return ""


# ============================================================
# Clean HTML
# ============================================================

def clean_html(html: str):
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["table", "sup", "style", "script"]):
        tag.decompose()

    text = soup.get_text(separator="\n")

    lines = [line.strip() for line in text.splitlines() if line.strip()]

    return "\n".join(lines)


# ============================================================
# Generate texts (NO CHUNKING)
# ============================================================

def generate_texts(category_name):
    titles = WIKIPEDIA_SOURCES[category_name]

    texts = []

    for title in titles:
        print(f"[INFO] Fetching: {title}")
        text = fetch_wikipedia_text(title)

        if text and len(text) > 500:  # ensure meaningful size
            texts.append(text)

    return texts


# ============================================================
# Save text files
# ============================================================

def save_text_variants(folder: Path, category: str, texts):
    styles = getSampleStyleSheet()

    for idx, text in enumerate(texts):

        # TXT
        with open(folder / f"{category}_{idx}.txt", "w", encoding="utf-8") as f:
            f.write(text)

        # Markdown
        with open(folder / f"{category}_{idx}.md", "w", encoding="utf-8") as f:
            f.write(f"# {category.capitalize()}\n\n{text}")

        # PDF
        pdf = SimpleDocTemplate(str(folder / f"{category}_{idx}.pdf"))
        pdf.build([Paragraph(text, styles["Normal"])])

        # DOCX
        doc = Document()
        doc.add_heading(category.capitalize(), 1)
        doc.add_paragraph(text)
        doc.save(folder / f"{category}_{idx}.docx")


# ============================================================
# Main builder
# ============================================================

def build_dataset():
    print("========================================")
    print("Building dataset (FULL ARTICLES, NO CHUNKING)")
    print("========================================")

    verify_coco()
    create_structure()

    with open(DATA_DIR / "annotations" / "instances_val2017.json") as f:
        coco = json.load(f)

    id_to_cat = {c["id"]: c["name"] for c in coco["categories"]}
    id_to_img = {img["id"]: img["file_name"] for img in coco["images"]}

    counters = {folder: 0 for folder in CATEGORY_MAPPING.values()}
    used_images = set()

    print("[INFO] Extracting images...")

    for ann in tqdm(coco["annotations"]):
        cat_name = id_to_cat[ann["category_id"]]

        if cat_name in CATEGORY_MAPPING:
            target = CATEGORY_MAPPING[cat_name]

            if counters[target] >= IMAGES_PER_CATEGORY:
                continue

            img_id = ann["image_id"]
            if img_id in used_images:
                continue

            filename = id_to_img[img_id]
            src = DATA_DIR / "val2017" / filename

            base_name = Path(filename).stem
            dst_base = OUTPUT_DIR / target / "images" / base_name

            resize_and_save_image(src, dst_base)

            counters[target] += 1
            used_images.add(img_id)

        if all(v >= IMAGES_PER_CATEGORY for v in counters.values()):
            break

    print("[OK] Image extraction completed.")

    print("[INFO] Generating full Wikipedia articles...")

    for category in CATEGORY_MAPPING.values():
        texts = generate_texts(category)
        save_text_variants(OUTPUT_DIR / category / "texts", category, texts)

    # Ground truth
    ground_truth = {
        category: {"expected_folder": category}
        for category in CATEGORY_MAPPING.values()
    }

    with open(OUTPUT_DIR / "ground_truth.json", "w", encoding="utf-8") as f:
        json.dump(ground_truth, f, indent=2)

    print("[SUCCESS] Dataset built.")
    print("✔ Full articles (no chunking)")
    print("✔ One document = one article")


# ============================================================

if __name__ == "__main__":
    build_dataset()